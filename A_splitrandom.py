"""Random and Unpatterned Variation

This module handles dictionary entries with random and unpatterned variations, 
ensuring accurate processing of complex English expressions within the parsing pipeline. 
It identifies, segments, constructs, and integrates new entries, enhancing the overall 
effectiveness of the package's data refinement process.


all subsequent python modules only work with two type of dictionary entries:

1. entries with a single English expression in the entry head
   e.g. 
   *an A for effort Fig. acknowledgement for having tried
    to do something, even if it was not successful. (*Typically:
    get ~; give someone ~.) _ The plan didn’t work, but
    I’ll give you an A for effort for trying.
2. entries with multiple English expressions in the entry head
    e.g.
    bail someone out of jail and bail someone out† 1. Lit. to
    deposit a sum of money that allows someone to get out of
    jail while waiting for a trial. _John was in jail. I had to go
    down to the police station to bail him out. _ I need some
    cash to bail out a friend! 2. Fig. to help someone who is
    having difficulties. _ When my brother went broke, I had
    to bail him out with a loan.

in both cases (as shown above), the English expressions are always located at the very 
beginning of the entry (entry head) and then followed by one or more senses/definitions, each
comes with one or more dictionary examples

over 90% of the entries in the book fall into either one of the types mentioned above.

a third type (not yet implemented) has multiple English expressions that can show up anywhere in
entry body, and those expressions has multiples senses that could be shared between one expression
and not the others
e.g.
give someone a lift 1. and give someone a ride Fig. to
provide transportation for someone. _ I’ve got to get into
town. Can you give me a lift? 2. Fig. to raise someone’s
spirits; to make a person feel better. _ It was a good conversation,
and her kind words really gave me a lift.

in the example above:
- 2 expressions - give someone a lift & give someone a ride
- only the first expression has two senses. the second expression only has one sense
- that means if we were to parse this entry using C_readit.py as if it was type 2 entry, 
  the results will be confusing and misleading to the end user
  {
    "phrase": "give someone a lift",
    "definition": "1. to provide transportation for someone. _ I’ve got to get into town. 
                   Can you give me a lift? 2. to raise someone’s spirits; 
                   to make a person feel better. _ It was a good conversation, 
                   and her kind words really gave me a lift.",
  }
  {
    "phrase": "give someone a ride",
    "definition": "1. to provide transportation for someone. _ I’ve got to get into town. 
                   Can you give me a lift? 2. to raise someone’s spirits; 
                   to make a person feel better. _ It was a good conversation, 
                   and her kind words really gave me a lift.",
  }


the code below will parse such entries and generate new dictionary entries that can be either
a type 1 or 2. and then process it normally using subsequent python modules

Thought process:
step 1: locate all entries in clean_output.docx that are type 3 and capture their respective ranges
(count = 128)
step 2: cut each entry into smaller pieces. each piece should contain an entry head 
(with one or more English expressions) with the subsequent senses.
using the example above, it would be broken down into two pieces
[
    give someone a lift 1.
]
[
    give someone a ride Fig. to
    provide transportation for someone. _ I’ve got to get into
    town. Can you give me a lift? 2. Fig. to raise someone’s
    spirits; to make a person feel better. _ It was a good conversation,
    and her kind words really gave me a lift.
]
step 3: construct new entries by rearranging alt/runs to give each entry the relevant 
definition/examples
rule to follow: the first entry head inherits all subsequent senses, the rest only gets
the single sense associated with it

using the example above, it will be converted to the following two entries
[
    give  someone  a lift 1.  Fig.  to provide transportation for 
    someone.  _  I’ve got to get into town. Can you give me a lift?  
    2.  Fig.  to raise someone’s spirits; to make a person feel better.
    _  It was a good conversation, and her kind words really gave me a lift.
]
[
    give  someone  a ride  Fig.  to provide transportation for someone. 
    _  I’ve got to get into town. Can you give me a lift? 
]

step 4: write the newly generated entries to clean_output.docx and capture their new ranges 
- (start, end) for each entry (265 new entries gets generated)

step 5: remove captured_ranges from ranges.pickle, and then add new_ranges to it.

"""

import pickle
import re

import docx
from docx.shared import Pt  # for run.font.size = Pt(r.font.size.pt)

from Z_module import copy_docx, runtype

# load entry ranges []
with open("files/ranges.pickle", "rb") as file:
    ranges = pickle.load(file)

# open up clean-output.docx
doc = docx.Document("files/clean-output.docx")
lines = doc.paragraphs

# step 1
captured_ranges = []
for s, e in ranges:

    entry_runs = []  # capture all entry runs

    for i in range(s, e + 1):
        runs = lines[i].runs

        for ri, run in enumerate(runs):
            entry_runs.append(run)

    for i, r in enumerate(entry_runs):
        # examine entry runs - look for '[1-9]\. and'
        try:
            if (
                r.bold
                and re.search(r"[1-9]\.", r.text.strip())
                and (
                    runtype(i, entry_runs[i + 1]) == "and"
                    or runtype(i, entry_runs[i + 2]) == "and"
                    # or statement for (3620, 3628), (74394, 74400)
                )
            ):
                captured_ranges.append((s, e))
                break
        except IndexError:
            continue

# save captured entries to a docx file - for better readability
copy_docx(captured_ranges, "entries_with_random_variation")


# step 2
entry_pieces = []  # [([[alt1], [alt2]], [[runs1], [runs2]]), (), (), ...]
for s, e in captured_ranges:
    entry_alt = []
    entry_runs = []

    tmp_alt = []
    tmp_runs = []

    split_alt = []
    split_runs = []

    for i in range(s, e + 1):
        runs = lines[i].runs

        for ri, run in enumerate(runs):
            entry_alt.append(runtype(ri, run, mode="MUL"))
            entry_runs.append(run)

        # keep the original formatting - add new line at the end of each paragraph
        entry_alt.append("new_line")
        entry_runs.append("new_line")

    # break up entry_alt & entry_runs into smaller lists where each is an independent entry
    for c, (a, r) in enumerate(zip(entry_alt, entry_runs)):
        if a == "new_line":
            tmp_alt.append(a)
            tmp_runs.append(r)
            continue
        try:
            if runtype(c, r) == "and" and (
                (
                    entry_runs[c - 1].bold
                    and re.search(r"[1-9]\.", entry_runs[c - 1].text.strip())
                )
                or (
                    # this condistion is for [42891, 42899]
                    runtype(c, entry_runs[c - 1]) == "term"
                    and entry_runs[c - 2].bold
                    and re.search(r"[1-9]\.", entry_runs[c - 2].text.strip())
                )
            ):
                split_alt.append(tmp_alt)
                split_runs.append(tmp_runs)
                tmp_alt = []
                tmp_runs = []
            else:
                tmp_alt.append(a)
                tmp_runs.append(r)  # capture run as is - not in text format
        except IndexError:
            continue

    split_alt.append(tmp_alt)
    split_runs.append(tmp_runs)

    # for c, (la, lr) in enumerate(zip(split_alt, split_runs), start=1):
    #     print(f"{c}\n{la}\n", "-" * 50, f"\n{lr}")
    # print("*" * 50)
    entry_pieces.append((split_alt, split_runs))


# step 3
def headless_entry(entry_alt, entry_runs):
    """return entry body (definition/examples) after removing entry head

    Args:
        entry_alt (list): list of entry alternatives
        entry_runs (list): list of entry runs
    """
    # find indexes for definition|example|term|article
    indexes = []
    items_to_find = ["definition", "term", "example", "article"]

    for item in items_to_find:
        try:
            index = entry_alt.index(item)
            indexes.append(index)
        except ValueError:
            # Handle the case where the item is not found in entry_alt
            pass

    # sort indexes and remove duplicates
    sorted_indexes = sorted(set(indexes))
    min_index = (
        sorted_indexes[0] if sorted_indexes[0] != 0 else sorted_indexes[1]
    )  # smallest index is the start of

    return entry_alt[min_index:], entry_runs[min_index:]


def single_sense(entry_alt, entry_runs):
    """onyl retain the first sense/definition in a an entry with multiple senses/definitions

    Args:
        entry_alt (list): list of alternitaves
        entry_runs (list): list of runs
    """
    # get new_sense index
    indexes = []
    try:
        index = entry_alt.index("new_sense")
        indexes.append(index)
    except ValueError:
        pass
    sorted_indexes = sorted(set(indexes))

    if sorted_indexes and sorted_indexes[0] != 0:
        return entry_alt[: sorted_indexes[0]], entry_runs[: sorted_indexes[0]]
    if len(sorted_indexes) > 1 and sorted_indexes[0] == 0:
        return entry_alt[: sorted_indexes[1]], entry_runs[: sorted_indexes[1]]
    return entry_alt, entry_runs


for alts, runs in entry_pieces:

    for indx in range(1, len(alts)):
        # add all senses to the very first entry
        additional_alt, additional_runs = headless_entry(alts[indx], runs[indx])
        alts[0].extend(additional_alt)
        runs[0].extend(additional_runs)
        # make sure all senses have a single definition - starting from index #1
        alts[indx], runs[indx] = single_sense(alts[indx], runs[indx])

# need to capture run.text in line 209 for this to work
# for alts, runs in entry_pieces:
#     print("Entry")
#     for i in range(len(alts)):
#         print("".join(runs[i]))
#         print("-" * 50)

#     print("\n", "*" * 50)


# step 4
# write newly constructed entries to clean-output.docx


def runs_to_lines(lst_of_runs):
    """break a list of run on "new_line"

    Args:
        lst_of_runs (list): list of items that contain the string "new_line"
        that should be used to break up the list.

    """
    lines = []
    tmp_line = []
    for indx, itm in enumerate(lst_of_runs):
        if indx == 0 and itm == "new_line":
            continue
        if indx == (len(lst_of_runs) - 1) and itm == "new_line":
            break
        if itm == "new_line":
            lines.append(tmp_line)
            tmp_line = []
        else:
            tmp_line.append(itm)

    if tmp_line:
        lines.append(tmp_line)

    return lines


# open up clean-output.docx
doc = docx.Document("files/clean-output.docx")
lines = doc.paragraphs

# i've hardcoded the total number of lines in the original file of clean-output.docx to avoid
# writing the new entries a again into the document with every new run to this module.
line_number = 89728

FIRST_TIME = True  # check if this was the first time we ran this module or not
if len(lines) - 1 > line_number:
    FIRST_TIME = False

new_ranges = []

for alts, runs in entry_pieces:

    for new_entry in runs:

        # break [new_entry] into lines e.g. [[line1], [line2], ...]
        new_entry_lines = runs_to_lines(new_entry)
        entry_line_numbers = []  # new entry line numbers in the docx file

        for line in new_entry_lines:
            paragraph = doc.add_paragraph()  # initiate blank line
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)

            for r in line:
                # add each run to the new paragraph/line in doc
                run = paragraph.add_run(r.text)
                # apply original run style
                if r.bold:
                    run.bold = True
                if r.italic:
                    run.italic = True
                run.font.name = r.font.name
                run.font.size = Pt(r.font.size.pt)
            # Increment line number counter after adding the paragraph
            line_number += 1
            entry_line_numbers.append(line_number)

        new_ranges.append(
            (entry_line_numbers[0], entry_line_numbers[-1])
        )  # store as tuple (start, end)

if FIRST_TIME:
    doc.save("files/clean-output.docx")

# step 5
with open("files/ranges.pickle", "rb") as file:
    ranges = pickle.load(file)

# remove captured_ranges from ranges.pickle
ranges = [r for r in ranges if r not in captured_ranges]
# add new_ranges
ranges.extend(item for item in new_ranges if item not in ranges)

with open("files/ranges.pickle", "wb") as file:
    pickle.dump(ranges, file)
