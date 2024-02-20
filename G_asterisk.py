""" Parse Verbs Associated with the Asterisk in Dictionary Entries

Description:
This script analyzes and updates dictionary entry data, focusing on entries starting with an asterisk (*).

Input:
- phrases.json
- clean-output.docx

Output:
- Updated phrases.json

Runtime:
- Generating the updated phrases.json: Completed in 5 seconds.

Usage:
Please run this script from the command line (CMD)

Example:
python G_tidyup.py
"""

import json
import re

import docx

from Z_module import CompactJSONEncoder


def asterisk_verbs(wip_alt, wip_runs, entry_range):
    """
    Analyzes and updates dictionary entry data for expressions starting with an asterisk (*).

    This function checks if a dictionary entry begins with an asterisk (*) by examining the 'runs' list of the entry.
    If an asterisk is found, it extracts and parses verbs associated with the asterisk, which serve as key points for later stages in identifying expressions in sentences. It then updates the 'alt' and 'runs' lists with this information.

    Args:
    wip_alt (list of str): The 'alt' list of the entry.
    wip_runs (list of str): The 'runs' list of the entry.
    entry_range (tuple): A tuple containing the start and end lines of the entry in 'clean-output.docx'.

    Note:
    This function should be called after creating the 'lines' variable, which includes the lines of the 'clean-output.docx' document, using the following:
    ```
    doc = docx.Document("files/clean-output.docx")
    lines = doc.paragraphs
    ```

    Returns:
    tuple: A tuple containing two lists:
        - Updated 'alt' list (list of str) with added 'verb' markers for verbs associated with the asterisk.
        - Updated 'runs' list (list of str) with extracted verbs at the beginning.

    If the entry does not begin with an asterisk, the function returns (None, None).

    Example:
    Given the following dictionary entry:
    *accustomed to someone or something and *accustomed
    to doing something used to someone or something; used to
    or in the habit of doing something. (*Typically: be ~;
    become ~; grow ~.) _ The children are accustomed
    to eating late in the evening.

    Input:
    wip_alt = ['constant', 'variable']
    wip_runs = ['accustomed to', 'someone or something']

    Expected Output:
    alt = ['verb', 'verb', 'verb', 'constant', 'variable']
    runs = ['be', 'become', 'grow', 'accustomed to', 'someone or something']

    This function will affect or alter 992 entries in the dictionary.
    """

    if wip_runs[0] == "*" or wip_runs[0].startswith("*"):
        verbs = list()
        entry_runs = list()
        entry_font = list()
        asterisk = False
        # capture verbs in bold that lies between '*Typically:' or '*Also:' and the next closing parenthesis ')'. an entry could have multiple verb forms.
        for l in range(entry_range[0], entry_range[1] + 1):
            runs = lines[l].runs
            for run in runs:
                if "*Typically:" in run.text or "*Also:" in run.text:
                    asterisk = True
                    continue
                if asterisk and entry_runs and ")" in run.text:
                    asterisk = False
                elif asterisk:
                    entry_runs.append(run)
                    entry_font.append(run.bold)

        for r, f in zip(entry_runs, entry_font):
            if f and "*Typically:" not in r.text and "*Also:" not in r.text:
                verbs.append(r.text)

        # clean up [verbs]
        verbs = [re.sub(r"[~;.)(†]", "", v).strip() for v in verbs]
        # remove duplicates
        verbs = list(dict.fromkeys(verbs))
        # remove empty strings
        verbs = list(filter(None, verbs))
        # remove the asterisk '*' from both 'alt' and 'runs'
        if wip_runs[0] == "*":
            del wip_alt[0]
            del wip_runs[0]
        elif wip_runs[0].startswith("*"):
            wip_runs[0] = wip_runs[0].replace("*", "")
        # add verbs at the beginning of 'alt' and 'runs'
        alt = ["verb"] * len(verbs) + wip_alt
        runs = verbs + wip_runs

        return alt, runs

    else:
        return None, None


def fix_into(wip_alt, wip_runs):
    """
    Fix dictionary entries that begin with an asterisk (*) and contain '[into]' in the verbs, while the first 'constant' in the 'runs' list starts with 'in'.

    This function is designed to correct dictionary entries that meet the following criteria:
    - The entry begins with an asterisk (*).
    - One or more verbs generated by the 'asterisk_verbs()' function include '[into]'.
    - The first 'constant' in the 'runs' list starts with 'in'.

    In such cases, the 'into' from the 'verb' can replace the 'in' in the 'constant', resulting in two versions of the entry:
    1. The original entry with only the verbs that correspond to 'in'.
    2. A duplicate entry with verbs that correspond to 'into' and 'in' replaced with 'into' in the first 'constant'.

    Args:
    wip_alt (list of str): The 'alt' list of the entry.
    wip_runs (list of str): The 'runs' list of the entry.

    Returns:
    tuple: A tuple containing four lists:
        - Updated 'alt' list for the original entry (list of str).
        - Updated 'runs' list for the original entry (list of str).
        - 'alt' list for the duplicate entry (list of str).
        - 'runs' list for the duplicate entry (list of str).

    If the entry does not meet the specified criteria, the function returns (None, None, None, None).

    Example:
    Given the following dictionary entry:
    *in the middle of nowhere Fig. in a very remote place.
    (*Typically: be ~; drive [into] ~; put someone or
    something [into] ~.) _ To get to my house, you have to
    drive into the middle of nowhere. _ We found a nice place
    to eat, but it’s out in the middle of nowhere.

    Input:
    wip_alt = ['verb', 'verb', 'verb', 'verb', 'constant']
    wip_runs = ['be', 'drive [into]', 'put', '[into]', 'in the middle of nowhere']

    Expected Output:
    fixed_alt = ['verb', 'constant']
    fixed_runs = ['be', 'in the middle of nowhere']
    duplicate_alt = ['verb', 'verb', 'constant']
    duplicate_runs = ['drive', 'put', 'into the middle of nowhere']

    This function will affect or alter 53 entries in the dictionary.
    """

    into = False
    in_constant = False
    for a, r in zip(wip_alt, wip_runs):
        if a == "verb" and "[into]" in r:
            into = True
        if a == "constant" and r.startswith("in"):
            in_constant = True

    if into and in_constant:
        # 1- fix current entry
        fixed_alt = list(wip_alt)
        fixed_runs = list(wip_runs)

        indexes = list()
        for i, (a, r) in enumerate(zip(fixed_alt, fixed_runs)):
            if a == "verb" and r == "[into]":
                indexes.extend([i - 1, i])
            elif a == "verb" and "[into]" in r:
                indexes.append(i)
        for x in sorted(indexes, reverse=True):
            del fixed_alt[x]
            del fixed_runs[x]

        # 2- create a duplicate entry
        duplicate_alt = list()
        duplicate_runs = list()

        for i, (a, r) in enumerate(zip(wip_alt, wip_runs)):
            if a == "verb" and r == "[into]":
                duplicate_alt.append("verb")
                duplicate_runs.append(wip_runs[i - 1])
            elif a == "verb" and "[into]" in r:
                duplicate_alt.append("verb")
                duplicate_runs.append(wip_runs[i].replace("[into]", "").strip())
        for i, (a, r) in enumerate(zip(wip_alt, wip_runs)):
            if a != "verb":
                duplicate_alt.append(wip_alt[i])
                duplicate_runs.append(wip_runs[i])
        # replace 'in' in the first constant with 'into'
        for i, (a, r) in enumerate(zip(duplicate_alt, duplicate_runs)):
            if a == "constant" and r.startswith("in"):
                duplicate_runs[i] = r.replace("in", "into", 1).strip()
                break

        return fixed_alt, fixed_runs, duplicate_alt, duplicate_runs

    else:
        return None, None, None, None


# open up clean-output.docx
doc = docx.Document("files/clean-output.docx")
lines = doc.paragraphs

# load the json file
with open("englishidioms/phrases.json", encoding="UTF-8") as f:
    data = json.load(f)

# create a json string
json_string = """
{
  "dictionary": [

  ]
}
"""

# load json to python from a string - use file as a placeholder for all entries
file = json.loads(json_string)

for entry in data["dictionary"]:
    asterisk_alt, asterisk_runs = asterisk_verbs(
        entry["alt"], entry["runs"], entry["range"]
    )

    if None not in (asterisk_alt, asterisk_runs):
        E, F, G, H = fix_into(asterisk_alt, asterisk_runs)

        if all(x is not None for x in (E, F, G, H)):
            data1 = {
                "range": entry["range"],
                "phrase": entry["phrase"],
                "phrase_html": entry["phrase_html"],
                "definition": entry["definition"],
                "definition_html": entry["definition_html"],
                "alt": E,
                "runs": F,
                "multiple": entry["multiple"],
                "duplicate": entry["definition"],
            }

            file["dictionary"].append(data1)

            data2 = {
                "range": entry["range"],
                "phrase": entry["phrase"],
                "phrase_html": entry["phrase_html"],
                "definition": entry["definition"],
                "definition_html": entry["definition_html"],
                "alt": G,
                "runs": H,
                "multiple": entry["multiple"],
                "duplicate": True,
            }

            file["dictionary"].append(data2)

        else:
            new_data = {
                "range": entry["range"],
                "phrase": entry["phrase"],
                "phrase_html": entry["phrase_html"],
                "definition": entry["definition"],
                "definition_html": entry["definition_html"],
                "alt": asterisk_alt,
                "runs": asterisk_runs,
                "multiple": entry["multiple"],
                "duplicate": entry["duplicate"],
            }

            file["dictionary"].append(new_data)

    else:
        new_data = {
            "range": entry["range"],
            "phrase": entry["phrase"],
            "phrase_html": entry["phrase_html"],
            "definition": entry["definition"],
            "definition_html": entry["definition_html"],
            "alt": entry["alt"],
            "runs": entry["runs"],
            "multiple": entry["multiple"],
            "duplicate": entry["duplicate"],
        }

        file["dictionary"].append(new_data)


# overwrite the file
with open("englishidioms/phrases.json", "w", encoding="UTF-8") as f:
    json.dump(file, f, indent=2, cls=CompactJSONEncoder, ensure_ascii=False)
