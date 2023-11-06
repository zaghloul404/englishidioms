"""
Utility Functions for Multiple Modules

This script contains a collection of utility functions that are shared and utilized by multiple modules within the project.
Grouping these functions in a single script enhances code modularity, making it easier to maintain and update the functions from a central location.

Feel free to extend this script with additional utility functions as needed to support various parts of the project.
"""

import sys
import os
import re
import pickle
import json
import docx
from tqdm import tqdm
from docx.shared import Pt  # for run.font.size = Pt(r.font.size.pt)
import mammoth  # to convert docx to html


class CompactJSONEncoder(json.JSONEncoder):
    """A JSON Encoder that puts small containers on single lines."""

    CONTAINER_TYPES = (list, tuple, dict)
    """Container datatypes include primitives or other containers."""

    MAX_WIDTH = 100
    """Maximum width of a container that might be put on a single line."""

    MAX_ITEMS = 10
    """Maximum number of items in container that might be put on single line."""

    def __init__(self, *args, **kwargs):
        # using this class without indentation is pointless
        if kwargs.get("indent") is None:
            kwargs["indent"] = 4
        super().__init__(*args, **kwargs)
        self.indentation_level = 0

    def encode(self, o):
        """Encode JSON object *o* with respect to single line lists."""
        if isinstance(o, (list, tuple)):
            return self._encode_list(o)
        if isinstance(o, dict):
            return self._encode_object(o)
        if isinstance(o, float):  # Use scientific notation for floats
            return format(o, "g")
        return json.dumps(
            o,
            skipkeys=self.skipkeys,
            ensure_ascii=self.ensure_ascii,
            check_circular=self.check_circular,
            allow_nan=self.allow_nan,
            sort_keys=self.sort_keys,
            indent=self.indent,
            separators=(self.item_separator, self.key_separator),
            default=self.default if hasattr(self, "default") else None,
        )

    def _encode_list(self, o):
        if self._put_on_single_line(o):
            return "[" + ", ".join(self.encode(el) for el in o) + "]"
        self.indentation_level += 1
        output = [self.indent_str + self.encode(el) for el in o]
        self.indentation_level -= 1
        return "[\n" + ",\n".join(output) + "\n" + self.indent_str + "]"

    def _encode_object(self, o):
        if not o:
            return "{}"
        if self._put_on_single_line(o):
            return (
                "{ "
                + ", ".join(
                    f"{self.encode(k)}: {self.encode(el)}" for k, el in o.items()
                )
                + " }"
            )
        self.indentation_level += 1
        output = [
            f"{self.indent_str}{json.dumps(k)}: {self.encode(v)}" for k, v in o.items()
        ]

        self.indentation_level -= 1
        return "{\n" + ",\n".join(output) + "\n" + self.indent_str + "}"

    def iterencode(self, o, **kwargs):
        """Required to also work with `json.dump`."""
        return self.encode(o)

    def _put_on_single_line(self, o):
        return (
            self._primitives_only(o)
            and len(o) <= self.MAX_ITEMS
            and len(str(o)) - 2 <= self.MAX_WIDTH
        )

    def _primitives_only(self, o: list | tuple | dict):
        if isinstance(o, (list, tuple)):
            return not any(isinstance(el, self.CONTAINER_TYPES) for el in o)
        elif isinstance(o, dict):
            return not any(isinstance(el, self.CONTAINER_TYPES) for el in o.values())

    @property
    def indent_str(self) -> str:
        if isinstance(self.indent, int):
            return " " * (self.indentation_level * self.indent)
        elif isinstance(self.indent, str):
            return self.indentation_level * self.indent
        else:
            raise ValueError(
                f"indent must either be of type int or str (is: {type(self.indent)})"
            )


def runtype(ri, run, mode=""):
    """
    Function Description:
    The runtype function examines the runs produced by python-docx and generates an alternative (a meaning) for each run based on its text style properties.
    It returns the alternative as a string.

    Function Parameters:
    - ri (int): Stands for run index. It represents the index of the run being examined.
    - run (str): The actual run, which is a text string extracted from 'clean-output.docx'.
    - mode (optional, str): An optional parameter to consider ['the', 'a', 'an'] as a term regardless of its index in the entry.
    This is primarily used to fix certain cases where alt, run is not being handled correctly. By default, mode is an empty string.

    Function Returns:

    - alternative (str): A string representing the alternative meaning or category of the run based on its text style properties.

    The runtype function takes a run and analyzes its text style properties to determine its meaning.
    It checks for various style properties, such as boldness, font name, and font size, to categorize the run into different alternatives.
    The mode parameter allows for special handling of certain cases where alt, run needs to be adjusted.

    Note:
    parameter mode affects the following 94 entries
    [[634, 642], [701, 706], [1695, 1698], [4872, 4886], [4978, 4981], [5209, 5213], [5239, 5246], [5323, 5325], [6023, 6028], [6394, 6400], [6556, 6561], [8503, 8508], [8670, 8679], [14612, 14619], [14620, 14628], [15271, 15274], [15586, 15591], [16272, 16275], [16399, 16404], [17238, 17248], [17267, 17269], [17270, 17273], [20502, 20508], [20845, 20849], [21608, 21618], [22796, 22802], [22860, 22864], [23392, 23395], [23883, 23888], [23981, 23990], [26601, 26606], [31062, 31068], [31263, 31266], [31442, 31448], [31549, 31556], [31799, 31803], [32188, 32192], [32607, 32610], [32631, 32634], [34303, 34305], [34598, 34604], [34623, 34628], [36078, 36081], [36082, 36087], [36869, 36873], [37143, 37148], [37627, 37631], [44409, 44419], [44424, 44428], [47005, 47012], [47042, 47044], [47345, 47349], [47540, 47548], [47575, 47583], [48171, 48174], [48305, 48308], [48318, 48321], [48322, 48328], [50327, 50332], [50664, 50673], [50808, 50812], [51641, 51647], [52986, 52989], [55055, 55059], [56938, 56942], [57356, 57360], [58966, 58974], [64653, 64658], [64659, 64663], [65400, 65402], [65408, 65412], [66320, 66323], [66388, 66391], [68832, 68835], [71819, 71825], [72045, 72048], [72390, 72397], [72613, 72618], [73210, 73214], [73849, 73856], [75139, 75141], [76109, 76111], [78661, 78668], [79184, 79187], [79952, 79957], [80210, 80214], [81854, 81859], [81904, 81907], [82682, 82688], [83441, 83446], [83984, 83986], [84390, 84394], [87046, 87049], [87970, 87978]]
    13 of those 94 entries are actually ['the', 'a', 'an'] at the very end of the line and in fact the beginning of a definition - will fix this later
    """

    terms = [
        "Cliché",
        "Euph.",
        "Fig.",
        "Inf.",
        "Lit.",
        "Prov.",
        "Rur.",
        "Sl.",
        "Euph",
        "Fig",
        "Inf",
        "Lit",
        "Prov",
        "Rur",
        "Sl",
        "Prov. Cliché",
        "Fig. Cliché",
        "Fig. Inf.",
        "Fig. Sl.",
        "Fig. Euph.",
        "Cliché Fig.",
        "Euph. Fig.",
        "Inf. Fig.",
        "Lit. Fig.",
        "Inf. Lit.",
    ]

    if (
        run.text.strip() == "*"
        and run.bold
        and run.font.name == "Formata-Medium"
        and run.font.size.pt == float("8.5")
    ):
        return "asterisk"
    elif (
        (ri == 0 or ri == 1)
        and run.text.strip().lower() in ["the", "a", "an"]
        and run.font.name == "Minion-Regular"
        and run.font.size.pt == float("8.5")
    ) or (
        mode == "MUL"
        and run.text.strip().lower() in ["the", "a", "an"]
        and run.font.name == "Minion-Regular"
        and run.font.size.pt == float("8.5")
    ):
        return "article"
    elif (
        run.bold
        and run.font.name == "Formata-Medium"
        and run.font.size.pt == float("8.5")
    ):
        return "constant"
    elif run.font.name == "Formata-Condensed" and run.font.size.pt == float("9.0"):
        return "variable"
    elif (
        (run.text.strip() in terms)
        and run.italic
        and (run.font.name == "Minion-Italic")
        and (run.font.size.pt == float("8.5"))
    ):
        return "term"
    elif (
        run.text.strip() == "and"
        and run.font.name == "Minion-Regular"
        and int(run.font.size.pt) == 7
    ):
        return "and"
    elif (run.font.name == "Minion-Regular" and run.font.size.pt == float("8.5")) or (
        run.font.name == "Formata-Regular" and run.font.size.pt == float("8.0")
    ):
        return "definition"
    elif (
        (run.text.strip() == "_")
        and (run.font.name == "MathematicalPi-Six")
        and (run.font.size.pt == float("8.5"))
    ):
        return "underscore"
    elif (
        (run.text.strip() not in terms)
        and run.italic
        and (run.font.name == "Minion-Italic")
        and (run.font.size.pt == float("8.5"))
    ):
        return "example"
    elif (
        run.bold
        and run.font.name == "Formata-Medium"
        and run.font.size.pt == float("4.5")
        and run.text.strip() == "†"
    ):
        return "dagger"
    else:
        return "None"


def cleanup(line_alt_list, line_runs_list):
    """
    Function Description:
    The cleanup function examines the line_alt_list and line_runs_list produced by python-docx and runtype(), and it removes any items that should not be part of an entry head.
    It returns modified line_alt_list and line_runs_list that represent an entry head.

    Function Parameters:
    - line_alt_list (list): A list containing alternative meanings for entry runs.
    - line_runs_list (list): A list containing runs (sections of text) for any entry.

    Function Returns:
    - line_alt_list (list): A modified list of alternative meanings representing an entry head.
    - line_runs_list (list): A modified list of runs representing an entry head.

    The cleanup function is designed to process and refine the provided lists to isolate the entry head from other content in a document.
    It performs several checks to remove specific elements that are not part of the entry head and returns the cleaned lists.
    """

    try:
        def_index = line_alt_list.index("definition")
        line_alt_list = line_alt_list[:def_index]
        line_runs_list = line_runs_list[:def_index]
    except:
        # only two entries have no definition, ignore them for now
        pass

    # remove all items in both lists from the 1st 'term' and beyond
    try:
        term_index = line_alt_list.index("term")
        line_alt_list = line_alt_list[:term_index]
        line_runs_list = line_runs_list[:term_index]
    except:
        # not all entries have a term
        pass

    # remove all items in both lists from the 1st 'example' and beyond
    try:
        example_index = line_alt_list.index("example")
        line_alt_list = line_alt_list[:example_index]
        line_runs_list = line_runs_list[:example_index]
    except:
        pass

    # remove '1.' if it was picked up at the very end of entry head - 1030 match
    if line_alt_list[-1] == "constant" and line_runs_list[-1].strip() == "1.":
        line_alt_list.pop(-1)
        line_runs_list.pop(-1)

    # remove 'article' if it was picked up at the very end of entry head - 13 match (all from multiple phrase entries)
    if line_alt_list[-1] == "article" and line_runs_list[-1].strip().lower() in [
        "the",
        "a",
        "an",
    ]:
        line_alt_list.pop(-1)
        line_runs_list.pop(-1)

    return line_alt_list, line_runs_list


def copy_docx(ranges_list, output_file):
    """
    a function to copy lines (within a given list of ranges) from clean-output.docx into output_file.docx
    created a function as I needed to use this code multiple times
    """

    # open up clean-output.docx
    doc = docx.Document("files/clean-output.docx")
    doc_lines = doc.paragraphs

    print(f"creating {output_file}.docx")
    new_doc = docx.Document()  # create a new document
    for s, e in tqdm(ranges_list):
        # add range
        paragraph_1 = new_doc.add_paragraph()  # initiate blank line
        paragraph_1.paragraph_format.space_before = Pt(1)
        paragraph_1.paragraph_format.space_after = Pt(1)
        add_range = paragraph_1.add_run(f"[{s}, {e}]")
        add_range.font.name = "Calibri"
        add_range.font.size = docx.shared.Pt(10)

        # add the actual entry lines
        for line in range(s, e + 1):
            paragraph_2 = (
                new_doc.add_paragraph()
            )  # add an empty paragraph/line to tmp_doc
            paragraph_2.paragraph_format.space_before = Pt(1)
            paragraph_2.paragraph_format.space_after = Pt(1)

            # break down the line I want to copy from 'clean-output.docx' into runs
            doc_runs = doc_lines[line].runs

            # copy runs from 'clean-output.docx' to new_doc and maintain the same style
            for i, r in enumerate(doc_runs):
                # add each run to the new paragraph/line in tmp_doc
                run = paragraph_2.add_run(r.text)

                # apply original run style
                if r.bold:
                    run.bold = True
                if r.italic:
                    run.italic = True
                run.font.name = r.font.name
                run.font.size = Pt(r.font.size.pt)

        # add a break after each entry
        paragraph_3 = new_doc.add_paragraph()
        paragraph_3.paragraph_format.space_before = Pt(1)
        paragraph_3.paragraph_format.space_after = Pt(1)
        break_line = paragraph_3.add_run(f'{"*" * 36}')
        break_line.font.name = "Courier New"
        break_line.font.size = docx.shared.Pt(10)

    new_doc.save(f"files/{output_file}.docx")  # save the docx file locally


def parse_entry(entry_range, cut_off, multiple_phrases=False, line_runs=[]):
    """
    Function Description:
    The parse_entry function extracts and processes an entry head and body from the "clean-output.docx" document, returning the phrase text, phrase HTML, and definition HTML for the entry.

    Function Parameters:
    - entry_range (tuple): A tuple (s, e) representing the range of lines to be copied from "clean-output.docx." The function will extract content from line s to line e.
    - cut_off (int): The index that marks the end of the entry head. All content before this index is considered part of the entry head, and content after it is part of the entry body.
    - multiple_phrases (bool, optional): A flag indicating whether the entry head contains multiple phrases. If True, the function will extract and process each phrase separately. Default is False.
    - line_runs (list, optional): A list of runs representing lines. This is used when multiple_phrases is True to extract individual phrases.

    Function Usage:
    1. The function checks whether the entry with the specified entry_range has been processed before by loading entries from a pickled file called "entry_details.pickle." If the entry has been processed previously, it returns the cached results.
    2. If the entry is not found in the cache, it opens the "clean-output.docx" document and extracts its content.
    3. The entry head content is saved to "entry_head.docx," and the entry body content is saved to "entry_body.docx."
    4. If multiple_phrases is True, the entry head is divided into individual phrases, and both phrase text and HTML are generated.
    5. For each phrase, temporary docx files ("phrase_0.docx," "phrase_1.docx," etc.) are created and processed, and the results are cached.
    6. HTML for the entry body is generated and processed, removing terms and making modifications to the HTML content.
    7. The temporarily created docx files are deleted, and the results are pickled and cached.
    8. The function returns phrase text, phrase HTML, and entry body HTML.

    Example Usage:
    ```
    entry_range = (5, 10)  # Specify the range of lines to extract
    cut_off = 7  # Index marking the end of the entry head
    multiple_phrases = True  # If entry head has multiple phrases
    line_runs = [...]  # List of runs representing lines (only needed when multiple_phrases is True)

    phrase_text, phrase_html, entry_body_html = parse_entry(entry_range, cut_off, multiple_phrases, line_runs)
    ```
    """

    #######################################################
    # check to see if this entry has been processed before
    if not os.path.exists("files/entry_details.pickle"):
        # create the file if this was the 1st script run
        entry_details = []
        with open("files/entry_details.pickle", "wb") as file:
            pickle.dump(entry_details, file)
    else:
        # load the file
        with open("files/entry_details.pickle", "rb") as file:
            entry_details = pickle.load(file)

    for entry in entry_details:
        if entry[0] == entry_range:
            return entry[1], entry[2], entry[3]
    #######################################################

    # open up clean-output.docx
    doc = docx.Document("files/clean-output.docx")
    lines = doc.paragraphs

    # create a temp docx document for entry head
    tmp_doc_1 = docx.Document()

    runs_count = 0  # counter

    # let's copy the entry head from clean-output.docx and save it to entry_head.docx
    for i in range(entry_range[0], entry_range[1] + 1):
        p = tmp_doc_1.add_paragraph()  # add an empty paragraph/line to tmp_doc

        # break down the line I want to copy from 'clean-output.docx' into runs
        doc_runs = lines[i].runs

        for r in doc_runs:
            # check if I should add run or not
            if runs_count < cut_off:
                # add run to the new paragraph/line in tmp_doc
                run = p.add_run(r.text)

                # apply original run style
                if r.bold:
                    run.bold = True
                if r.italic:
                    run.italic = True
                run.font.name = r.font.name
                run.font.size = Pt(r.font.size.pt)

                # increment runs_count
                runs_count += 1
            else:
                break  # break from the inner loop
        else:
            # Continue if the inner loop wasn't broken.
            continue
        # Inner loop was broken, break the outer.
        break

    # save the docx file locally
    tmp_doc_1.save("entry_head.docx")

    if multiple_phrases:
        ## step #1 - create phrase_text - a list of strings, each is an individual phrase in the entry head
        #######################
        phrase_text = []
        for lr in line_runs:
            pt = " ".join(lr)
            pt1 = pt.replace("1.", "").replace("* ", "*")
            phrase_text.append(pt1.strip())

        ## step #2 - create phrase_html -  a list of strings, each is an HTML code block that represents a phrase in the entry head
        #######################
        phrase_html = []
        tmp_docx_files = []

        # open up entry_head.docx
        eh_doc = docx.Document("entry_head.docx")
        eh_lines = eh_doc.paragraphs

        # count how many runs are there in entry_head.docx & add all entry head runs to a list
        total_runs_count = 0
        entry_head_runs = []
        for l in eh_lines:
            eh_runs = l.runs
            for r in eh_runs:
                total_runs_count += 1
                entry_head_runs.append(r)
            # add a single space ' ' after each line
            total_runs_count += 1
            entry_head_runs.append(" ")

        # read entry head and save each individual phrase to a new docx file
        run_counter = 0
        done = []  # runs that I've captured already
        for i in range(20):
            # create a tmp docx file to host a single phrase
            sp_doc = docx.Document()  # single phrase document
            # add an empty paragraph/line to tmp_doc
            sp_p = sp_doc.add_paragraph()

            # loop over the full entry head and save phrases to tmp_doc one by one
            for c, r in enumerate(entry_head_runs):
                if c == run_counter and r == " ":
                    sp_p.add_run(r)
                    run_counter += 1
                    done.append((c, r, " "))

                elif c == run_counter and runtype(c, r) != "and" and ";" not in r.text:
                    # add run to the new paragraph/line in tmp_doc
                    tmp_run = sp_p.add_run(r.text)

                    # apply original run style
                    if r.bold:
                        tmp_run.bold = True
                    if r.italic:
                        tmp_run.italic = True
                    tmp_run.font.name = r.font.name
                    tmp_run.font.size = Pt(r.font.size.pt)

                    run_counter += 1  # increment counter
                    done.append((c, r, r.text))

                elif c == run_counter and ";" in r.text:
                    all_parts = r.text.split(";")
                    # in case r.text.startswith(';') or r.text.endswith(';') it will add an empty string to all_parts - fix it
                    # if len(all_parts[-1]) == 0: all_parts.pop(-1)
                    if len(all_parts[-1]) == 0:
                        all_parts[-1] = ";"
                    # if len(all_parts[0])  == 0: all_parts.pop(0)
                    if len(all_parts[0]) == 0:
                        all_parts[0] = ";"

                    # select the part that is in line
                    for pc, p in enumerate(all_parts, start=1):
                        if p not in [x[2] for x in done if x[0] == c]:
                            run_text = p
                            break

                    # add part to the new paragraph/line in tmp_doc
                    tmp_run = sp_p.add_run(run_text)

                    # apply original run style
                    if r.bold:
                        tmp_run.bold = True
                    if r.italic:
                        tmp_run.italic = True
                    tmp_run.font.name = r.font.name
                    tmp_run.font.size = Pt(r.font.size.pt)

                    done.append((c, r, run_text))

                    # only increment run_counter if we have been through all run parts
                    if run_text == all_parts[-1] and pc == len(all_parts):
                        run_counter += 1
                        continue
                    else:
                        break

                elif c == run_counter and (
                    runtype(c, r) == "and" or r.text.strip() == ";"
                ):
                    run_counter += 1
                    done.append((c, r, r.text))

                    break

            # save the docx file locally
            sp_doc.save(f"phrase_{i}.docx")
            tmp_docx_files.append(f"phrase_{i}.docx")

            # check if we have reached the end of the file
            if run_counter == total_runs_count:
                break

        # convert all docx files to HTML code and save it to [phrase_html]
        for f in tmp_docx_files:
            with open(f, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                entry_head_html = result.value  # The generated HTML

            # clean up the HTML output
            phrase_html.append(
                entry_head_html.replace("1.", "")
                .replace(" </p>", "</p>")
                .replace(" </strong></p>", "</strong></p>")
                .replace("</p><p>", " ")
                .replace("<strong>;</strong>", "")
            )

            # delete the file as it's no longer needed
            try:
                os.remove(f)
            except:
                pass

        # check to make sure everything is working as it should be
        if len(phrase_text) != len(phrase_html):
            print("*" * 25)
            print("something went wrong".upper())
            print("entry_range")
            print(entry_range)
            print("line_runs")
            print(line_runs)
            print("phrase_text")
            print(phrase_text)
            print("len(phrase_text)")
            print(len(phrase_text))
            print("phrase_html")
            print(phrase_html)
            print("len(phrase_html)")
            print(len(phrase_html))
            print("*" * 25)
            sys.exit()
    else:
        # grab phrase text from tmp_doc
        tmp_a_phrase_text = []
        for l in tmp_doc_1.paragraphs:
            tmp_a_phrase_text.append(l.text.strip())
        tmp_b_phrase_text = " ".join(tmp_a_phrase_text)
        tmp_c_phrase_text = tmp_b_phrase_text.replace("1.", "")
        phrase_text = tmp_c_phrase_text.strip()

        # get HTML of the entry head
        with open("entry_head.docx", "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            entry_head_html = result.value  # The generated HTML

        # clean up the HTML output
        phrase_html = (
            entry_head_html.replace("1.", "")
            .replace(" </p>", "</p>")
            .replace(" </strong></p>", "</strong></p>")
            .replace("</p><p>", " ")
        )

    # create a temp docx document for entry body
    tmp_doc_2 = docx.Document()

    runs_count = 0  # reset counter

    for i in range(entry_range[0], entry_range[1] + 1):
        p = tmp_doc_2.add_paragraph()  # add an empty paragraph/line to tmp_doc

        # break down the line I want to copy from 'clean-output.docx' into runs
        doc_runs = lines[i].runs

        for r in doc_runs:
            # check if I should add run or not
            if runs_count >= cut_off:
                # add run to the new paragraph/line in tmp_doc
                run = p.add_run(r.text)

                # apply original run style
                if r.bold:
                    run.bold = True
                if r.italic:
                    run.italic = True
                run.font.name = r.font.name
                run.font.size = Pt(r.font.size.pt)

            # increment runs_count
            runs_count += 1

    # save the docx file locally
    tmp_doc_2.save("entry_body.docx")

    # generate HTML for entry body
    with open("entry_body.docx", "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        generated_html = result.value  # The generated HTML

    # remove terms from definitions
    terms = [
        "Cliché",
        "Euph.",
        "Fig.",
        "Inf.",
        "Lit.",
        "Prov.",
        "Rur.",
        "Sl.",
        "Euph",
        "Fig",
        "Inf",
        "Lit",
        "Prov",
        "Rur",
        "Sl",
        "Prov. Cliché",
        "Fig. Cliché",
        "Fig. Inf.",
        "Fig. Sl.",
        "Fig. Euph.",
        "Cliché Fig.",
        "Euph. Fig.",
        "Inf. Fig.",
        "Lit. Fig.",
        "Inf. Lit.",
    ]
    for term in terms:
        generated_html = generated_html.replace(
            f"{term} ", ""
        )  # remove terms and training whitespace
        generated_html = generated_html.replace(term, "")

    # let's make some modifications to the HTML
    # https://stackoverflow.com/questions/59072514/efficiently-make-many-multiple-substitutions-in-a-string/59072515#59072515

    replace = {
        "</p><p>": " ",
        "_": "<br>_",
        "<strong>2. </strong>": "<br><strong>2. </strong>",
        "<strong>3. </strong>": "<br><strong>3. </strong>",
        "<strong>4. </strong>": "<br><strong>4. </strong>",
        "<strong>5. </strong>": "<br><strong>5. </strong>",
        "<strong>6. </strong>": "<br><strong>6. </strong>",
        "<strong>7. </strong>": "<br><strong>7. </strong>",
        "<strong>8. </strong>": "<br><strong>8. </strong>",
        "<strong>9. </strong>": "<br><strong>9. </strong>",
        "<strong>10. </strong>": "<br><strong>10. </strong>",
    }

    rep = dict((re.escape(k), v) for k, v in replace.items())
    pattern = re.compile("|".join(rep.keys()))
    entry_body_html = pattern.sub(lambda m: rep[re.escape(m.group(0))], generated_html)

    if "<strong>2." in entry_body_html and "<strong>1." not in entry_body_html:
        # add "<strong>1. </strong>" to the beginning of the string
        entry_body_html = entry_body_html.replace(
            "<p>", "<p><strong>1. </strong>", 1
        )  # only replace the 1st occurrence

    # remove all (See also)
    entry_body_html = re.sub(
        r"\(See also.+?\)\.\)|\(See also.+?\)", "", entry_body_html
    )

    # delete the temporarily created docx files
    try:
        os.remove("entry_head.docx")
        os.remove("entry_body.docx")
    except:
        pass

    # pickle results
    results = [entry_range, phrase_text, phrase_html, entry_body_html]
    entry_details.append(results)
    with open("files/entry_details.pickle", "wb") as file:
        pickle.dump(entry_details, file)

    # return results
    return phrase_text, phrase_html, entry_body_html


if __name__ == "__main__":
    pass
